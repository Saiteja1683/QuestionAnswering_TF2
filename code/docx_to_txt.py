import docx
import os
import pandas as pd
import numpy as np

import nltk 
from nltk.corpus import stopwords 
from nltk.tokenize import word_tokenize, sent_tokenize 

stop_words = set(stopwords.words('english')) 
print(stop_words)
csv_data = pd.read_csv('./data/TrainingTestSet.csv')


questions = ['what is the payble amount of rent per month or monthly rent?',
			'which is the starting(begin) date or contract made date or commencing date  or executed day of agreement from ?',
			'which is the ending date or duration or validation of agreement termination to ?',
			'how many days of prior notice(contract) period for termination(remain) of agreement ?',
			'who is owner or lessor or management or Procall or Landlord ?',
			'who is tenant or lessee or resident or user ?']

max_pragraph_length = 100




main_dirc = './data/Training_data/'  # training doc files
output_txt_dirc = './data/training_txt_data/'

try :
	os.mkdir(output_txt_dirc)
except:
	pass
context_txt = open(output_txt_dirc+'contexts.txt','w')
question_txt = open(output_txt_dirc+'questions.txt','w')
ans_txt = open(output_txt_dirc+'answers.txt','w')

files = os.listdir(main_dirc)

for i in files:
	print(i)
	amount = (csv_data.loc[csv_data['File Name'] == i[:-9],'Aggrement Value']).iloc[0]
	start_date = (csv_data.loc[csv_data['File Name'] == i[:-9],'Aggrement Start Date']).iloc[0]
	end_date = (csv_data.loc[csv_data['File Name'] == i[:-9],'Aggrement End Date']).iloc[0]
	notice_days = (csv_data.loc[csv_data['File Name'] == i[:-9],'Renewal Notice (Days)']).iloc[0]
	owner = (csv_data.loc[csv_data['File Name'] == i[:-9],'Party One']).iloc[0]
	tenant = (csv_data.loc[csv_data['File Name'] == i[:-9],'Party Two']).iloc[0]

	print(nltk.pos_tag([str(amount),str(start_date),str(end_date),str(notice_days),str(owner),str(tenant)]))

	doc = docx.Document(main_dirc+i)
	processed_paras = []
	print(len(doc.paragraphs),'total length..................!!!!')

	for n in range(len(doc.paragraphs)):
		paragraphs = doc.paragraphs[n].text
		if len(paragraphs)!=0:
			wordsList = nltk.word_tokenize(paragraphs) 
			wordsList = [w for w in wordsList if not w in stop_words]  
			tagged = nltk.pos_tag(wordsList)
			lis_nnp = [i[0] for i in tagged if i[-1]=="NNP" or i[-1]=="NP" or i[-1]=="CD" or i[-1]=="RB" or i[-1]=="NNS" or i[-1]=="JJ" or i[-1]=="NN"]
			
			if len(lis_nnp)!=0: 
				processed_paras.append(paragraphs)
	
	print(len(processed_paras),'processed............!!!')

	for feat in range(0,5):
		print(feat)
		for para in processed_paras:
			context_txt.write(para)	
		context_txt.write('\n')
	for que in questions:
		question_txt.write(que+'\n')
	ans_txt.write(str(amount)+'\n')
	ans_txt.write(str(start_date)+'\n')
	ans_txt.write(str(end_date)+'\n')
	ans_txt.write(str(notice_days)+'\n')
	ans_txt.write(str(owner)+'\n')
	ans_txt.write(str(tenant)+'\n')
		
	
	

